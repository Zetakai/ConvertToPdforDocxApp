import sys
import os
import sqlite3
import pandas as pd
import markdown2
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from PyQt6.QtWidgets import (
    QApplication, QWidget, QLabel, QPushButton, QVBoxLayout, QMessageBox, QFileDialog, QLineEdit, QCheckBox, QGroupBox, QHBoxLayout, QProgressBar, QComboBox
)
from PyQt6.QtGui import QPixmap
from PyQt6.QtCore import (Qt,QTimer)
from PIL import Image
import platform
from reportlab.lib.pagesizes import letter, landscape
from reportlab.pdfgen import canvas
from reportlab.platypus import Table, TableStyle
from reportlab.lib import colors
import pandas as pd
import pdfplumber
from docx.shared import Inches

# Initialize the SQLite database
def init_db():
    conn = sqlite3.connect("users.db")
    cursor = conn.cursor()
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT UNIQUE,
            password TEXT
        )
    """)
    conn.commit()
    conn.close()

# Register window
class RegisterWindow(QWidget):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Register")
        self.resize(350, 250)

        # Input Group
        input_group = QGroupBox("Create a New Account")
        self.username_label = QLabel("Username:")
        self.username_input = QLineEdit()
        
        self.password_label = QLabel("Password:")
        self.password_input = QLineEdit()
        self.password_input.setEchoMode(QLineEdit.EchoMode.Password)

        self.show_password_checkbox = QCheckBox("Show Password")
        self.show_password_checkbox.stateChanged.connect(self.toggle_password_visibility)

        # Layout for input fields
        input_layout = QVBoxLayout()
        input_layout.addWidget(self.username_label)
        input_layout.addWidget(self.username_input)
        input_layout.addWidget(self.password_label)
        input_layout.addWidget(self.password_input)
        input_layout.addWidget(self.show_password_checkbox)
        input_group.setLayout(input_layout)

        # Register button
        self.register_button = QPushButton("Register")
        self.register_button.clicked.connect(self.register_user)
        
        # Main layout
        main_layout = QVBoxLayout()
        main_layout.addWidget(input_group)
        main_layout.addWidget(self.register_button)
        main_layout.setSpacing(15)
        self.setLayout(main_layout)
    
    def toggle_password_visibility(self):
        if self.show_password_checkbox.isChecked():
            self.password_input.setEchoMode(QLineEdit.EchoMode.Normal)
        else:
            self.password_input.setEchoMode(QLineEdit.EchoMode.Password)
    
    def register_user(self):
        username = self.username_input.text()
        password = self.password_input.text()
        
        if not username or not password:
            QMessageBox.warning(self, "Error", "Both fields are required.")
            return

        conn = sqlite3.connect("users.db")
        cursor = conn.cursor()
        try:
            cursor.execute("INSERT INTO users (username, password) VALUES (?, ?)", (username, password))
            conn.commit()
            QMessageBox.information(self, "Success", "Registration successful!")
            self.clear_inputs()
            self.close()
        except sqlite3.IntegrityError:
            QMessageBox.warning(self, "Error", "Username already exists.")
        finally:
            conn.close()
    
    def clear_inputs(self):
        self.username_input.clear()
        self.password_input.clear()
        self.show_password_checkbox.setChecked(False)

def open_file(pdf_path, event=None):
    if platform.system() == "Windows":
        os.startfile(pdf_path)
    elif platform.system() == "Darwin":  # macOS
        os.system(f'open "{pdf_path}"')
    elif platform.system() == "Linux":
        os.system(f'xdg-open "{pdf_path}"')

# Menu window with file conversion functionality
class MenuWindow(QWidget):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("File Conversion to PDF or DOCX")
        self.resize(500, 300)
        
        # File selection group
        file_selection_group = QGroupBox("File Selection")
        self.supported_files_label = QLabel("Supported file types: DOCX, CSV, TXT, XLSX, XLS, PNG, JPG, JPEG, PDF")
        self.file_label = QLabel("No file selected.")
        self.select_button = QPushButton("Select File")
        self.select_button.clicked.connect(self.select_file)

        file_selection_layout = QVBoxLayout()
        file_selection_layout.addWidget(self.supported_files_label)
        file_selection_layout.addWidget(self.file_label)
        file_selection_layout.addWidget(self.select_button)
        file_selection_group.setLayout(file_selection_layout)

        # Conversion format selection
        self.format_combo = QComboBox()
        self.format_combo.addItem("Convert to PDF")
        self.format_combo.addItem("Convert to DOCX")

        # Conversion and logout group
        action_group = QGroupBox("Actions")
        self.convert_button = QPushButton("Convert")
        self.convert_button.clicked.connect(self.start_conversion)

        self.logout_button = QPushButton("Logout")
        self.logout_button.clicked.connect(self.logout)

        action_layout = QHBoxLayout()
        action_layout.addWidget(self.format_combo)
        action_layout.addWidget(self.convert_button)
        action_layout.addWidget(self.logout_button)
        action_group.setLayout(action_layout)

        # Progress bar for conversion
        self.progress_bar = QProgressBar()
        self.progress_bar.setValue(0)

        # Main layout
        main_layout = QVBoxLayout()
        main_layout.addWidget(file_selection_group)
        main_layout.addWidget(action_group)
        main_layout.addWidget(self.progress_bar)
        
        self.setLayout(main_layout)
        self.file_path = None
        self.timer = QTimer()
        self.timer.timeout.connect(self.update_progress)

    def select_file(self):
        file_dialog = QFileDialog()
        file_path, _ = file_dialog.getOpenFileName(
            self, "Select File", "", "Files (*.docx *.csv *.txt *.xlsx *.xls *.png *.jpg *.jpeg *.pdf)"
        )
        if file_path:
            self.file_path = file_path
            self.file_label.setText(f"Selected: {file_path}")
            self.progress_bar.setValue(0)

    def start_conversion(self):
        if not self.file_path:
            QMessageBox.warning(self, "Warning", "Please select a file first.")
            return
        
        self.progress_bar.setValue(0)
        self.timer.start(100)  # Start timer for progress simulation

    def update_progress(self):
        # Simulate progress
        value = self.progress_bar.value()
        if value < 100:
            self.progress_bar.setValue(value + 10)
        else:
            self.timer.stop()
            # Check selected format and call the appropriate conversion function
            if self.format_combo.currentText() == "Convert to PDF":
                self.convert_to_pdf()
            elif self.format_combo.currentText() == "Convert to DOCX":
                self.convert_to_docx()

    def convert_to_pdf(self):
        try:
            if self.file_path.endswith('.docx'):
                self.convert_docx_to_pdf(self.file_path)
            elif self.file_path.endswith('.csv'):
                self.convert_csv_to_pdf(self.file_path)
            elif self.file_path.endswith('.txt'):
                self.convert_txt_to_pdf(self.file_path)
            elif self.file_path.endswith('.xlsx'):
                self.convert_excel_to_pdf(self.file_path)
            elif self.file_path.endswith('.xls'):
                self.convert_excel_to_pdf(self.file_path)
            elif self.file_path.endswith(('.png', '.jpg', '.jpeg')):
                self.convert_image_to_pdf(self.file_path)
            else:
                QMessageBox.warning(self, "Warning", "Unsupported file type.")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to convert file to PDF: {e}")

    def convert_to_docx(self):
        if not self.file_path:  # Ensure file_path is set
            QMessageBox.warning(self, "Error", "No file selected for conversion.")
            return

        # Define the output DOCX file path
        docx_path = self.file_path.rsplit('.', 1)[0] + '.docx'

        try:
            doc = Document()

            # Handle .csv files
            if self.file_path.endswith('.csv'):
                df = pd.read_csv(self.file_path)
                doc.add_paragraph(", ".join(df.columns))
                for _, row in df.iterrows():
                    doc.add_paragraph(", ".join(str(value) for value in row))

            # Handle .txt files
            elif self.file_path.endswith('.txt'):
                with open(self.file_path, 'r', encoding='utf-8') as file:
                    for line in file:
                        doc.add_paragraph(line.strip())

            # Handle .xlsx and .xls files
            elif self.file_path.endswith('.xlsx') or self.file_path.endswith('.xls'):
                df = pd.read_excel(self.file_path)
                doc.add_paragraph(", ".join(df.columns))
                for _, row in df.iterrows():
                    doc.add_paragraph(", ".join(str(value) for value in row))

            # Handle .png, .jpg, .jpeg files
            elif self.file_path.endswith(('.png', '.jpg', '.jpeg')):
                doc.add_paragraph("Image:")
                doc.add_picture(self.file_path, width=Inches(5))  # Add image with fixed width

            # Handle .pdf files
            elif self.file_path.endswith('.pdf'):
                with pdfplumber.open(self.file_path) as pdf:
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text:
                            doc.add_paragraph(text)

            else:
                QMessageBox.warning(self, "Error", "Unsupported file type for DOCX conversion.")
                return

            # Save the created DOCX file
            doc.save(docx_path)
            QMessageBox.information(self, "Success", f"File converted to DOCX: {docx_path}")
            open_file(docx_path)
            
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to convert file to DOCX: {e}")

    def convert_docx_to_pdf(self, file_path):
        doc = Document(file_path)
        pdf_path = file_path.replace('.docx', '.pdf')
        c = canvas.Canvas(pdf_path, pagesize=letter)
        text_object = c.beginText(50, 750)
        for para in doc.paragraphs:
            text_object.textLine(para.text)
        c.drawText(text_object)
        c.save()
        QMessageBox.information(self, "Success", f"DOCX converted to PDF: {pdf_path}")
        open_file(pdf_path)

    def convert_csv_to_pdf(self, file_path):
        df = pd.read_csv(file_path)
        pdf_path = file_path.replace('.csv', '.pdf')
        c = canvas.Canvas(pdf_path, pagesize=letter)
        text_object = c.beginText(50, 750)
        text_object.textLine(" | ".join(df.columns))
        text_object.textLine("-" * 50)
        for _, row in df.iterrows():
            text_object.textLine(" | ".join(str(value) for value in row))
        c.drawText(text_object)
        c.save()
        QMessageBox.information(self, "Success", f"CSV converted to PDF: {pdf_path}")
        open_file(pdf_path)

    def convert_txt_to_pdf(self, file_path):
        pdf_path = file_path.replace('.txt', '.pdf')
        c = canvas.Canvas(pdf_path, pagesize=letter)
        with open(file_path, "r") as file:
            text = file.read()
            text_object = c.beginText(50, 750)
            text_object.setFont("Helvetica", 12)
            text_object.textLines(text)
            c.drawText(text_object)
        c.save()
        QMessageBox.information(self, "Success", f"TXT converted to PDF: {pdf_path}")
        open_file(pdf_path)

    def convert_excel_to_pdf(self, file_path):
            # Check if the file is .xls or .xlsx and load it into a DataFrame
        if file_path.endswith('.xls') or file_path.endswith('.xlsx'):
            df = pd.read_excel(file_path)
        else:
            QMessageBox.warning(self, "Error", "Unsupported file type. Please select an Excel file (.xls or .xlsx).")
            return
        
        # Define the output PDF file path
        pdf_path = file_path.replace('.xls', '.pdf').replace('.xlsx', '.pdf')

        # Set up the PDF canvas
        c = canvas.Canvas(pdf_path, pagesize=landscape(letter))
        width, height = landscape(letter)

        # Convert DataFrame to a list of lists for Table
        data = [df.columns.to_list()] + df.values.tolist()

        # Create a Table with the data
        table = Table(data)
        
        # Style the table
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),  # Header row background
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),  # Header row text color
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),  # Center align all cells
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),  # Header row font
            ('FONTSIZE', (0, 0), (-1, -1), 10),  # Font size for all cells
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),  # Padding for header row
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),  # Background for other rows
            ('GRID', (0, 0), (-1, -1), 1, colors.black)  # Grid lines for all cells
        ]))

        # Calculate the table size and position it on the canvas
        table.wrapOn(c, width, height)
        table.drawOn(c, 30, height - 50 - len(data) * 15)  # Adjust position based on data length

        # Save the canvas
        c.save()

        QMessageBox.information(self, "Success", f"Excel converted to PDF: {pdf_path}")
        open_file(pdf_path)

    def convert_image_to_pdf(self, file_path):
        image = Image.open(file_path)
        pdf_path = file_path.rsplit('.', 1)[0] + '.pdf'
        image.convert('RGB').save(pdf_path, "PDF")
        QMessageBox.information(self, "Success", f"Image converted to PDF: {pdf_path}")
        open_file(pdf_path)


    def logout(self):
        self.close()
        login_window.show()  # Show the login window again after logout

# Login window
class LoginWindow(QWidget):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Login")
        self.resize(350, 250)
        
        self.image_label = QLabel()
        self.image_label.setPixmap(QPixmap("file.png").scaled(100, 100, Qt.AspectRatioMode.KeepAspectRatio))
        self.image_label.setAlignment(Qt.AlignmentFlag.AlignCenter)

        # Group box for login form
        login_group = QGroupBox("Login to Your Account")
        
        self.username_label = QLabel("Username:")
        self.username_input = QLineEdit()

        self.password_label = QLabel("Password:")
        self.password_input = QLineEdit()
        self.password_input.setEchoMode(QLineEdit.EchoMode.Password)

        self.show_password_checkbox = QCheckBox("Show Password")
        self.show_password_checkbox.stateChanged.connect(self.toggle_password_visibility)

        # Layout for input fields
        login_layout = QVBoxLayout()
        login_layout.addWidget(self.username_label)
        login_layout.addWidget(self.username_input)
        login_layout.addWidget(self.password_label)
        login_layout.addWidget(self.password_input)
        login_layout.addWidget(self.show_password_checkbox)
        login_group.setLayout(login_layout)

        # Buttons for login and register
        self.login_button = QPushButton("Login")
        self.login_button.clicked.connect(self.login_user)

        self.register_button = QPushButton("Register")
        self.register_button.clicked.connect(self.open_register_window)

        # Button layout
        button_layout = QHBoxLayout()
        button_layout.addWidget(self.login_button)
        button_layout.addWidget(self.register_button)

        # Main layout
        main_layout = QVBoxLayout()
        main_layout.addWidget(self.image_label)
        main_layout.addWidget(login_group)
        main_layout.addLayout(button_layout)
        main_layout.setSpacing(15)
        self.setLayout(main_layout)
    
    def toggle_password_visibility(self):
        if self.show_password_checkbox.isChecked():
            self.password_input.setEchoMode(QLineEdit.EchoMode.Normal)
        else:
            self.password_input.setEchoMode(QLineEdit.EchoMode.Password)
    
    def login_user(self):
        username = self.username_input.text()
        password = self.password_input.text()
        
        conn = sqlite3.connect("users.db")
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM users WHERE username = ? AND password = ?", (username, password))
        user = cursor.fetchone()
        conn.close()
        
        if user:
            QMessageBox.information(self, "Success", "Login successful!")
            self.clear_inputs()
            self.open_menu_window()
            self.close()
        else:
            QMessageBox.warning(self, "Error", "Invalid username or password.")
    
    def clear_inputs(self):
        self.username_input.clear()
        self.password_input.clear()
        self.show_password_checkbox.setChecked(False)
    
    def open_register_window(self):
        self.register_window = RegisterWindow()
        self.register_window.show()
    
    def open_menu_window(self):
        self.menu_window = MenuWindow()
        self.menu_window.show()

# Main application
app = QApplication(sys.argv)
init_db()  # Initialize the database

login_window = LoginWindow()
login_window.show()

sys.exit(app.exec())
